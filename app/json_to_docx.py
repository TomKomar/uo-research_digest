import os
import json
from docx import Document
from docx.shared import Mm, Pt
from dotenv import load_dotenv
import requests
load_dotenv()


def read_json(file_path):
    """Reads a JSON file and returns the data."""
    with open(file_path, 'r') as file:
        data = json.load(file)
    return data

def fix_json(text_to_fix, api_key):
    print('Fixing JSON')
    headers = {
        'Content-Type': 'application/json',
        'Authorization': f'Bearer {api_key}'
    }
    data = {
        "model": "gpt-4o-mini",
        "response_format": {"type": "json_object"},
        "messages": [
            {"role": "system", "content": "You are a helpful assistant that fixes JSON formatting of text that fails direct parsing. Respond with fixed JSON."},
            {"role": "user", "content": "Fix following text so it can be parsed as JSON: \n"+text_to_fix}
        ],
        "max_tokens": 1500,
        "n": 1,
        "stop": None,
        "temperature": 0.0
    }
    response = requests.post('https://api.openai.com/v1/chat/completions', headers=headers, json=data)
    return response.json()['choices'][0]['message']['content'].strip()


def create_word_document(data, output_path, openai_api_key):
    """Creates a Word document from the provided data."""
    doc = Document()

    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(12.7)
    section.bottom_margin = Mm(12.7)
    section.left_margin = Mm(12.7)
    section.right_margin = Mm(12.7)

    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(9)

    style = doc.styles['Heading 1']
    font = style.font
    font.size = Pt(11)

    style = doc.styles['Heading 2']
    font = style.font
    font.size = Pt(10)

    style = doc.styles['Heading 3']
    font = style.font
    font.size = Pt(10)

    page = 0
    for key, value in data.items():
        print(f'Page: {key}')
        page += 1
        if page > 1:
            doc.add_page_break()
        doc.add_heading(key.replace('.txt', ''), level=1)
        details = value
        try:
            details = eval(details)
        except:
            details = eval(fix_json(details, openai_api_key))
        for section, content in details.items():
            doc.add_heading(section.replace("_", " ").capitalize(), level=2)
            if isinstance(content, str):
                doc.add_paragraph(content)
            elif isinstance(content, list):
                list_text = "; ".join([f"({i+1}) {item}" for i, item in enumerate(content)])
                doc.add_paragraph(list_text)
            elif isinstance(content, dict):
                for sub_key, sub_value in content.items():
                    doc.add_heading(sub_key.replace("_", " ").capitalize(), level=3)
                    if isinstance(sub_value, str):
                        doc.add_paragraph(sub_value)
                    elif isinstance(sub_value, list):
                        list_text = "; ".join([f"({i+1}) {item}" for i, item in enumerate(sub_value)])
                        doc.add_paragraph(list_text)

    doc.save(output_path)
    print(f"Document saved to {output_path}")

def main():
    json_file_path = r'c:\Users\User\Documents\papers\20250115\20250115.json'
    output_doc_path = r'c:\Users\User\Documents\papers\20250115\20250115.docx'

    data = read_json(json_file_path)
    create_word_document(data, output_doc_path, os.environ['openai_api_key'])


if __name__ == "__main__":
    main()